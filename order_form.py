from docx import Document as d
from docx.enum.text import WD_ALIGN_PARAGRAPH
import assitent_info
import streamlit as st
doc = d('C:\\Users\\novikov.rn\\Desktop\\VECTORST\\template.docx')
def fulfil_temp(cblank):
    # Таблица шапки
    doc.tables[0].rows[0].cells[1].text = cblank['orderer']
    doc.tables[0].rows[0].cells[3].text = 'общепромышленное'
    doc.tables[0].rows[1].cells[1].text = cblank['object']
    doc.tables[0].rows[2].cells[1].text = cblank['manager']
    doc.tables[0].rows[2].cells[3].text = cblank['developer']
    doc.tables[0].autofit
    st.write(doc.tables[0].style)
    doc.tables[2].rows[1].cells[1].text = f"На основании {int(cblank['glycol'])}% {cblank['glycol type'][:-1]}я" if cblank['glycol'] else assitent_info.pure_water
    doc.tables[2].rows[1].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.tables[2].rows[3].cells[1].text = str(cblank['consumption']).replace(".",",")
    doc.tables[2].rows[3].cells[1].paragraphs[0].alignment=WD_ALIGN_PARAGRAPH.CENTER
    doc.save('C:\\Users\\novikov.rn\\Desktop\\VECTORST\\test.docx')
    doc.tables[2].autofit